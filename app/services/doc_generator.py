"""
Document generator — produces a professional bilingual .docx from
summary.json, summary_zh.json, and action_items.json.
"""
import json
import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

logger = logging.getLogger(__name__)

# ── Colour palette ────────────────────────────────────────────────────────────
COLOUR_TITLE   = RGBColor(0x1A, 0x1A, 0x2E)   # near-black navy
COLOUR_HEADING = RGBColor(0x16, 0x21, 0x3E)   # dark blue
COLOUR_ACCENT  = RGBColor(0x0F, 0x3E, 0x6F)   # medium blue  (key concept label)
COLOUR_META    = RGBColor(0x55, 0x55, 0x55)   # grey for metadata line
COLOUR_FORMULA = RGBColor(0x2E, 0x4A, 0x1E)   # dark green for formulas


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_colour: str) -> None:
    """Fill a table cell background (used for action-item table header)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def _add_page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type())


def docx_break_type():
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    br = _OE("w:br")
    br.set(_qn("w:type"), "page")
    return br   # used as: run._r.append(br)


def _page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    br = _OE("w:br")
    br.set(_qn("w:type"), "page")
    run._r.append(br)


def _heading_para(doc: Document, text: str, size_pt: int, bold: bool = True,
                  colour: RGBColor = COLOUR_HEADING, space_before: int = 12,
                  space_after: int = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.color.rgb = colour


def _bullet(doc: Document, text: str, size_pt: int = 11, indent_cm: float = 0.5) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(indent_cm / 2.54)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size_pt)


def _key_concept_row(doc: Document, term: str, definition: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    label = p.add_run(f"{term}: ")
    label.bold = True
    label.font.size = Pt(10.5)
    label.font.color.rgb = COLOUR_ACCENT
    defn = p.add_run(definition)
    defn.font.size = Pt(10.5)


def _formula_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(10)
    run.font.color.rgb = COLOUR_FORMULA


def _divider(doc: Document) -> None:
    """Thin horizontal rule using a bottom-border paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _action_items_table(doc: Document, items: list, zh: bool = False) -> None:
    if not items:
        p = doc.add_paragraph()
        p.add_run("No action items found." if not zh else "未发现待办事项。").italic = True
        return

    col_widths = [Inches(1.1), Inches(3.6), Inches(1.2), Inches(0.9)]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Header row
    headers_en = ["Type", "Description", "Due Date", "Urgency"]
    headers_zh = ["类型", "描述", "截止日期", "优先级"]
    headers = headers_zh if zh else headers_en
    hdr_cells = table.rows[0].cells
    for i, (cell, hdr) in enumerate(zip(hdr_cells, headers)):
        cell.width = col_widths[i]
        _set_cell_bg(cell, "1A1A2E")
        run = cell.paragraphs[0].add_run(hdr)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for item in items:
        row_cells = table.add_row().cells
        desc = item.get("description_zh", item["description"]) if zh else item["description"]
        due  = item.get("due_date") or ("N/A" if not zh else "无")
        row_cells[0].text = item.get("type", "")
        row_cells[1].text = desc
        row_cells[2].text = due
        row_cells[3].text = item.get("urgency", "").upper()
        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()   # breathing room after table


# ── Section builders ──────────────────────────────────────────────────────────

def _build_english_section(doc: Document, summary: dict, action_items: list,
                            meta: dict) -> None:
    title = f"Lecture Notes: {summary.get('lecture_title', 'Untitled')}"
    _heading_para(doc, title, size_pt=18, colour=COLOUR_TITLE,
                  space_before=0, space_after=6)

    # Metadata line
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    meta_text = (
        f"Generated: {meta['date']}  |  "
        f"Duration: {meta['duration_min']:.0f} min  |  "
        f"Words: {meta['word_count']:,}  |  "
        f"Topics: {len(summary.get('topics', []))}"
    )
    run = p.add_run(meta_text)
    run.font.size = Pt(9)
    run.font.color.rgb = COLOUR_META
    run.italic = True

    _divider(doc)

    for topic in summary.get("topics", []):
        _heading_para(doc, topic["heading"], size_pt=13, space_before=14, space_after=4)

        for pt in topic.get("summary", []):
            _bullet(doc, pt)

        concepts = topic.get("key_concepts", [])
        if concepts:
            _heading_para(doc, "Key Concepts", size_pt=10, bold=True,
                          colour=COLOUR_ACCENT, space_before=6, space_after=2)
            for kc in concepts:
                _key_concept_row(doc, kc.get("term", ""), kc.get("definition", ""))

        formulas = topic.get("formulas", [])
        if formulas:
            _heading_para(doc, "Formulas", size_pt=10, bold=True,
                          colour=COLOUR_FORMULA, space_before=6, space_after=2)
            for f in formulas:
                _formula_para(doc, f)

        _divider(doc)

    # Action items
    _heading_para(doc, "Action Items & Announcements", size_pt=14,
                  colour=COLOUR_TITLE, space_before=14, space_after=6)
    _action_items_table(doc, action_items, zh=False)


def _build_mandarin_section(doc: Document, summary_zh: dict,
                             action_items_zh: list, meta: dict) -> None:
    title_zh = summary_zh.get("lecture_title_zh") or summary_zh.get("lecture_title", "")
    _heading_para(doc, f"讲座总结：{title_zh}", size_pt=18, colour=COLOUR_TITLE,
                  space_before=0, space_after=6)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    meta_text = (
        f"生成时间：{meta['date']}  |  "
        f"时长：{meta['duration_min']:.0f} 分钟  |  "
        f"字数：{meta['word_count']:,}  |  "
        f"主题数：{len(summary_zh.get('topics', []))}"
    )
    run = p.add_run(meta_text)
    run.font.size = Pt(9)
    run.font.color.rgb = COLOUR_META
    run.italic = True

    _divider(doc)

    for topic in summary_zh.get("topics", []):
        heading_zh = topic.get("heading_zh") or topic.get("heading", "")
        _heading_para(doc, heading_zh, size_pt=13, space_before=14, space_after=4)

        for pt in topic.get("summary_zh") or topic.get("summary", []):
            _bullet(doc, pt)

        concepts = topic.get("key_concepts", [])
        if concepts:
            _heading_para(doc, "关键概念", size_pt=10, bold=True,
                          colour=COLOUR_ACCENT, space_before=6, space_after=2)
            for kc in concepts:
                term_label = kc.get("term", "")   # keep term in English
                defn = kc.get("definition_zh") or kc.get("definition", "")
                _key_concept_row(doc, term_label, defn)

        formulas = topic.get("formulas", [])
        if formulas:
            _heading_para(doc, "公式", size_pt=10, bold=True,
                          colour=COLOUR_FORMULA, space_before=6, space_after=2)
            for f in formulas:
                _formula_para(doc, f)

        _divider(doc)

    _heading_para(doc, "待办事项与通知", size_pt=14, colour=COLOUR_TITLE,
                  space_before=14, space_after=6)
    _action_items_table(doc, action_items_zh, zh=True)


# ── Public API ────────────────────────────────────────────────────────────────

def generate_docx(session_id: str) -> Path:
    """
    Build a bilingual .docx from summary.json / summary_zh.json /
    action_items.json in the session's data directory.

    Saves to outputs/{session_id}/lecture_notes.docx and returns that path.
    """
    from app.config import get_settings

    settings = get_settings()
    session_dir  = settings.session_data_dir(session_id)
    outputs_dir  = settings.session_outputs_dir(session_id)

    # Load data
    summary      = json.loads((session_dir / "summary.json").read_text(encoding="utf-8"))
    summary_zh   = json.loads((session_dir / "summary_zh.json").read_text(encoding="utf-8"))
    action_items = json.loads((session_dir / "action_items.json").read_text(encoding="utf-8"))
    action_items_zh = json.loads((session_dir / "action_items_zh.json").read_text(encoding="utf-8"))

    # Derive metadata from transcript if available
    duration_min = 0.0
    word_count   = 0
    transcript_json = session_dir / "transcript_corrected.json"
    if transcript_json.exists():
        td = json.loads(transcript_json.read_text(encoding="utf-8"))
        duration_min = td.get("duration_minutes", 0.0)
        word_count   = td.get("word_count", 0)

    meta = {
        "date": datetime.now().strftime("%d %B %Y"),
        "duration_min": duration_min,
        "word_count": word_count,
    }

    doc = Document()

    # Page margins (2 cm all sides)
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # Default paragraph font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    _build_english_section(doc, summary, action_items, meta)
    _page_break(doc)
    _build_mandarin_section(doc, summary_zh, action_items_zh, meta)

    output_path = outputs_dir / "lecture_notes.docx"
    doc.save(str(output_path))
    logger.info("[%s] Saved lecture_notes.docx (%d KB)", session_id,
                output_path.stat().st_size // 1024)
    return output_path
